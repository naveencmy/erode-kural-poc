"""Content Extraction Engine for Multi-Format Government Documents.

Handles PDF, Excel, CSV, Image scans, Word documents, and Plain text
with unified structure extraction and metadata profiling.
"""

import csv
import json
import logging
import mimetypes
import os
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

import pandas as pd

import config
from pipeline.ocr_engine import IndicOCREngine

logger = logging.getLogger("DocumentContentExtractor")

# Optional python-docx import
try:
    import docx
except ImportError:
    docx = None

# Optional PyMuPDF
try:
    import pymupdf as fitz
except ImportError:
    fitz = None


def detect_file_type(file_path: Union[str, Path]) -> str:
    """Detect file type via robust extension and MIME mapping supporting any format."""
    path = Path(file_path)
    ext = path.suffix.lower().lstrip(".")

    EXT_MAP = {
        "pdf": "pdf",
        "xlsx": "xlsx",
        "xls": "xlsx",
        "xlsm": "xlsx",
        "xlsb": "xlsx",
        "ods": "xlsx",
        "csv": "csv",
        "tsv": "csv",
        "png": "png",
        "jpg": "jpg",
        "jpeg": "jpg",
        "tiff": "png",
        "tif": "png",
        "webp": "png",
        "bmp": "png",
        "gif": "png",
        "svg": "txt",
        "docx": "docx",
        "doc": "docx",
        "rtf": "txt",
        "odt": "docx",
        "txt": "txt",
        "md": "txt",
        "markdown": "txt",
        "json": "txt",
        "xml": "txt",
        "html": "txt",
        "htm": "txt",
        "log": "txt",
        "yaml": "txt",
        "yml": "txt",
        "sql": "txt",
        "eml": "eml",
        "msg": "txt",
        "pptx": "pptx",
        "ppt": "pptx",
    }

    if ext in EXT_MAP:
        return EXT_MAP[ext]

    guessed_type, _ = mimetypes.guess_type(str(path))
    if guessed_type:
        MIME_MAP = {
            "application/pdf": "pdf",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
            "application/vnd.ms-excel": "xlsx",
            "application/vnd.oasis.opendocument.spreadsheet": "xlsx",
            "text/csv": "csv",
            "text/tab-separated-values": "csv",
            "image/png": "png",
            "image/jpeg": "jpg",
            "image/tiff": "png",
            "image/webp": "png",
            "image/bmp": "png",
            "image/gif": "png",
            "image/svg+xml": "txt",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
            "application/msword": "docx",
            "application/vnd.oasis.opendocument.text": "docx",
            "application/rtf": "txt",
            "text/plain": "txt",
            "text/markdown": "txt",
            "application/json": "txt",
            "application/xml": "txt",
            "text/xml": "txt",
            "text/html": "txt",
            "application/vnd.openxmlformats-officedocument.presentationml.presentation": "pptx",
        }
        if guessed_type in MIME_MAP:
            return MIME_MAP[guessed_type]

    return "unknown"


class ContentExtractor:
    """Extracts unified textual, tabular, and block structure from diverse document formats."""

    def __init__(self, ocr_engine: Optional[IndicOCREngine] = None):
        self.ocr_engine = ocr_engine or IndicOCREngine()

    def extract(self, file_path: Union[str, Path], file_type: Optional[str] = None) -> Dict[str, Any]:
        """Route to appropriate extractor and return unified structure."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {path}")

        ftype = file_type or detect_file_type(path)

        extractors = {
            "pdf": self._extract_pdf,
            "xlsx": self._extract_excel,
            "csv": self._extract_csv,
            "png": self._extract_image,
            "jpg": self._extract_image,
            "docx": self._extract_docx,
            "txt": self._extract_text,
            "pptx": self._extract_pptx,
            "eml": self._extract_text,
        }

        handler = extractors.get(ftype, self._extract_unknown)
        result = handler(path)
        result["file_type"] = ftype
        result["file_name"] = path.name
        result["file_size_bytes"] = path.stat().st_size
        return result

    def _extract_pptx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PowerPoint presentations or fallback."""
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            slides_text = []
            with zipfile.ZipFile(file_path, "r") as z:
                for filename in z.namelist():
                    if filename.startswith("ppt/slides/slide") and filename.endswith(".xml"):
                        tree = ET.fromstring(z.read(filename))
                        texts = [node.text for node in tree.iter() if node.text]
                        if texts:
                            slides_text.append(" ".join(texts))
            full_text = "\n\n".join(slides_text)
            if full_text:
                return {
                    "text": full_text,
                    "tables": [],
                    "blocks": [],
                    "page_count": max(1, len(slides_text)),
                    "has_images": False,
                    "amount_columns": [],
                }
        except Exception:
            pass
        return self._extract_unknown(file_path)


    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """PDF structure extraction with page text, Indic OCR, and table detection."""
        all_text = []
        all_tables = []
        all_blocks = []
        page_count = 0

        # Attempt PyMuPDF direct text and table extraction first
        if fitz is not None:
            try:
                doc = fitz.open(str(file_path))
                page_count = len(doc)
                for page_idx, page in enumerate(doc, start=1):
                    p_text = page.get_text() or ""
                    
                    # Extract tables if available in PyMuPDF
                    try:
                        tabs = page.find_tables()
                        for tab in tabs:
                            df_tab = tab.to_pandas()
                            if not df_tab.empty:
                                all_tables.append({
                                    "page": page_idx,
                                    "name": f"Page_{page_idx}_Table",
                                    "columns": [str(c) for c in df_tab.columns],
                                    "row_count": len(df_tab),
                                    "sample_rows": df_tab.head(5).to_dict(orient="records"),
                                    "contains_amounts": any("தொகை" in str(c) or "amount" in str(c).lower() or "₹" in str(c) for c in df_tab.columns),
                                })
                    except Exception:
                        pass

                    # Extract blocks
                    try:
                        raw_blocks = page.get_text("blocks")
                        for b in raw_blocks:
                            if len(b) >= 5 and str(b[4]).strip():
                                all_blocks.append({
                                    "page": page_idx,
                                    "bbox": [b[0], b[1], b[2], b[3]],
                                    "text": str(b[4]).strip(),
                                })
                    except Exception:
                        pass

                    all_text.append(f"--- [Page {page_idx}] ---\n{p_text.strip()}")
            except Exception as e:
                logger.warning(f"PyMuPDF extraction note: {e}")

        # Fallback / augment with Indic OCR if no text layer found
        if not "".join(all_text).strip():
            logger.info(f"PDF text layer empty. Processing with Indic OCR: {file_path}")
            ocr_results = self.ocr_engine.process_document(file_path, source_id=file_path.stem)
            page_count = len(ocr_results)
            for item in ocr_results:
                page_no = item.get("page_number", 1)
                text = item.get("full_text", "")
                all_text.append(f"--- [Page {page_no}] ---\n{text}")
                for blk in item.get("blocks", []):
                    all_blocks.append({
                        "page": page_no,
                        "bbox": blk.get("bbox", [0, 0, 800, 600]),
                        "text": blk.get("text", ""),
                        "confidence": blk.get("confidence", 0.9),
                    })

        full_text = "\n\n".join(all_text)
        return {
            "text": full_text,
            "tables": all_tables,
            "blocks": all_blocks,
            "page_count": max(1, page_count),
            "has_images": len(all_blocks) > 0,
            "amount_columns": [],
        }

    def _extract_excel(self, file_path: Path) -> Dict[str, Any]:
        """Excel structure extraction with multi-sheet detection, column profiling, and amounts."""
        xl = pd.ExcelFile(str(file_path))
        all_tables = []
        all_columns = []
        amount_columns = []
        text_snippets = []

        for sheet_name in xl.sheet_names[:5]:  # Top 5 sheets
            df = xl.parse(sheet_name)
            if df.empty:
                continue

            sheet_cols = []
            for col in df.columns:
                col_name = str(col).strip()
                dtype_str = str(df[col].dtype)
                col_info = {
                    "name": col_name,
                    "dtype": dtype_str,
                    "null_count": int(df[col].isnull().sum()),
                    "sample_values": [str(v) for v in df[col].dropna().head(5).tolist()],
                    "unique_count": int(df[col].nunique()),
                }
                sheet_cols.append(col_info)
                all_columns.append(col_info)

                # Check if currency / amount column
                col_lower = col_name.lower()
                is_amount = any(k in col_lower for k in ["amount", "தொகை", "ஒதுக்கீடு", "செலவு", "budget", "cost", "₹", "rs", "lakhs", "crores", "ரூபாய்"])
                if not is_amount and pd.api.types.is_numeric_dtype(df[col]):
                    max_val = df[col].max()
                    if pd.notnull(max_val) and max_val > 1000:
                        is_amount = True
                if is_amount:
                    amount_columns.append(col_name)

            all_tables.append({
                "name": sheet_name,
                "columns": [str(c) for c in df.columns],
                "row_count": len(df),
                "sample_rows": df.head(10).to_dict(orient="records"),
                "contains_amounts": len(amount_columns) > 0,
            })

            text_snippets.append(f"--- Sheet: {sheet_name} (Rows: {len(df)}) ---\n" + df.head(100).to_string(index=False))

        return {
            "text": "\n\n".join(text_snippets),
            "tables": all_tables,
            "columns": all_columns,
            "amount_columns": list(set(amount_columns)),
            "page_count": len(xl.sheet_names),
            "has_images": False,
            "blocks": [],
        }

    def _extract_csv(self, file_path: Path) -> Dict[str, Any]:
        """CSV structure extraction with schema profiling."""
        try:
            df = pd.read_csv(str(file_path), encoding="utf-8", nrows=10000)
        except UnicodeDecodeError:
            df = pd.read_csv(str(file_path), encoding="latin1", nrows=10000)

        columns = []
        amount_columns = []
        for col in df.columns:
            col_name = str(col).strip()
            col_info = {
                "name": col_name,
                "dtype": str(df[col].dtype),
                "null_count": int(df[col].isnull().sum()),
                "sample_values": [str(v) for v in df[col].dropna().head(5).tolist()],
                "unique_count": int(df[col].nunique()),
            }
            columns.append(col_info)

            col_lower = col_name.lower()
            if any(k in col_lower for k in ["amount", "தொகை", "ஒதுக்கீடு", "செலவு", "budget", "₹", "rs", "cost"]):
                amount_columns.append(col_name)
            elif pd.api.types.is_numeric_dtype(df[col]):
                max_val = df[col].max()
                if pd.notnull(max_val) and max_val > 1000:
                    amount_columns.append(col_name)

        return {
            "text": df.head(150).to_string(index=False),
            "tables": [{
                "name": file_path.stem,
                "columns": [str(c) for c in df.columns],
                "row_count": len(df),
                "sample_rows": df.head(10).to_dict(orient="records"),
                "contains_amounts": len(amount_columns) > 0,
            }],
            "columns": columns,
            "amount_columns": list(set(amount_columns)),
            "page_count": 1,
            "has_images": False,
            "blocks": [],
        }

    def _extract_image(self, file_path: Path) -> Dict[str, Any]:
        """Image scan extraction via Indic OCR."""
        ocr_result = self.ocr_engine.process_image(str(file_path), source_id=file_path.stem, page_number=1)
        full_text = ocr_result.get("full_text", "")
        blocks = ocr_result.get("blocks", [])
        return {
            "text": full_text,
            "tables": [],
            "blocks": blocks,
            "page_count": 1,
            "has_images": True,
            "amount_columns": [],
        }

    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """Word document extraction with paragraph and table parsing."""
        all_text = []
        all_tables = []

        if docx is not None:
            try:
                doc = docx.Document(str(file_path))
                for p in doc.paragraphs:
                    if p.text.strip():
                        all_text.append(p.text.strip())

                for t_idx, tab in enumerate(doc.tables, start=1):
                    rows_data = []
                    headers = []
                    for r_idx, row in enumerate(tab.rows):
                        cells = [c.text.strip() for c in row.cells]
                        if r_idx == 0:
                            headers = cells
                        else:
                            rows_data.append(dict(zip(headers, cells)))
                    if headers:
                        all_tables.append({
                            "name": f"Table_{t_idx}",
                            "columns": headers,
                            "row_count": len(rows_data),
                            "sample_rows": rows_data[:5],
                            "contains_amounts": any("தொகை" in h or "amount" in h.lower() or "₹" in h for h in headers),
                        })
            except Exception as e:
                logger.warning(f"python-docx error: {e}")

        full_text = "\n\n".join(all_text)
        return {
            "text": full_text,
            "tables": all_tables,
            "blocks": [],
            "page_count": max(1, len(all_text) // 20),
            "has_images": False,
            "amount_columns": [],
        }

    def _extract_text(self, file_path: Path) -> Dict[str, Any]:
        """Plain text / note extraction."""
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="latin1", errors="replace")

        return {
            "text": text,
            "tables": [],
            "blocks": [],
            "page_count": 1,
            "has_images": False,
            "amount_columns": [],
        }

    def _extract_unknown(self, file_path: Path) -> Dict[str, Any]:
        """Fallback extractor for unclassified binary or mixed formats."""
        try:
            text = file_path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            text = f"[Binary file {file_path.name}]"

        return {
            "text": text[:5000],
            "tables": [],
            "blocks": [],
            "page_count": 1,
            "has_images": False,
            "amount_columns": [],
        }
