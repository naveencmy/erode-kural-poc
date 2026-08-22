"""DOCX and PDF Exporter for Official Content Documents (Bilingual Tamil & English).

Generates formatted Microsoft Word (.docx) and PDF files with
authentic Tamil Nadu Government styling for official document export.
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import config
from modules.official_content.generator import detect_language

logger = logging.getLogger("OfficialContentExporter")

# TN Government colors
TN_NAVY = RGBColor(0x1A, 0x3A, 0x5C)
TN_GOLD = RGBColor(0xC8, 0xA9, 0x51)


def _extract_clean_body_paragraphs_and_footer(
    body_text: str,
    default_footer: str = "வெளியீடு செய்தி மக்கள் தொடர்பு அலுவலர், ஈரோடு மாவட்டம்."
) -> Tuple[List[str], str]:
    """Extract clean body paragraphs and separate/deduplicate issuing footer line."""
    if not body_text:
        return [], default_footer

    paragraphs = [p.strip() for p in body_text.split("\n\n") if p.strip()]
    cleaned = []
    detected_footer = None

    for p in paragraphs:
        lines = [l.strip() for l in p.split("\n") if l.strip()]
        kept_lines = []
        for line in lines:
            line_no_num = re.sub(r'^\d+[\.\)]\s*', '', line).strip()
            is_short = len(line_no_num) < 140

            if is_short and (
                re.match(r"^செ\.வெ\.எண்\s*[-–:]", line_no_num) or
                re.match(r"^சுற்றறிக்கை\s*எண்\s*[-–:]", line_no_num) or
                re.match(r"^குறிப்பாணை\s*எண்\s*[-–:]", line_no_num) or
                re.match(r"^எண்:\s*வே/", line_no_num) or
                re.match(r"^நாள்\s*[-–:]", line_no_num) or
                re.match(r"^press release no\s*[-–:]", line_no_num, re.IGNORECASE) or
                re.match(r"^circular no\s*[-–:]", line_no_num, re.IGNORECASE) or
                re.match(r"^memorandum no\s*[-–:]", line_no_num, re.IGNORECASE) or
                re.match(r"^roc\. no\s*[-–:]", line_no_num, re.IGNORECASE) or
                re.match(r"^date\s*[-–:]", line_no_num, re.IGNORECASE) or
                line_no_num in ["----", "---", "----------------", "<><><>"] or
                line_no_num == "அவர்களின் செய்திக்குறிப்பு-" or
                line_no_num == "அவர்களின் சுற்றறிக்கை-" or
                line_no_num == "அவர்களின் அலுவலகக் குறிப்பாணை-" or
                line_no_num == "ஈரோடு மாவட்ட ஆட்சித்தலைவர் திரு.ச.கந்தசாமி இ.ஆ.ப.," or
                line_no_num.startswith("--------------------------------") or
                line_no_num.startswith("================================")
            ):
                continue

            if is_short and (
                line_no_num.startswith("வெளியீடு") or
                "செய்தி மக்கள் தொடர்பு அலுவலர்" in line_no_num or
                line_no_num.startswith("இப்படிக்கு") or
                line_no_num.lower().startswith("issued by") or
                "public relations officer" in line_no_num.lower() or
                line_no_num.lower().startswith("by order")
            ):
                detected_footer = line_no_num
                continue

            kept_lines.append(line)

        if kept_lines:
            cleaned.append("\n".join(kept_lines))

    final_footer = default_footer
    if detected_footer:
        norm_det = re.sub(r'[\s\-_,.:;]', '', detected_footer)
        norm_def = re.sub(r'[\s\-_,.:;]', '', default_footer)
        if norm_det != norm_def:
            final_footer = detected_footer

    return cleaned, final_footer


def _extract_clean_body_paragraphs(body_text: str) -> List[str]:
    paras, _ = _extract_clean_body_paragraphs_and_footer(body_text)
    return paras


def _add_header(doc: Document, template_title: str, ref_number: str, date_str: str, officer_id: str, is_english: bool = False):
    """Add government-styled header to the document."""
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top_title = "GOVERNMENT OF TAMIL NADU — DISTRICT COLLECTORATE\n" if is_english else "தமிழ்நாடு அரசு — மாவட்ட ஆட்சியர் அலுவலகம்\n"
    run = header_para.add_run(top_title)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = TN_NAVY

    dist_title = "ERODE DISTRICT\n" if is_english else "ஈரோடு மாவட்டம்\n"
    run2 = header_para.add_run(dist_title)
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = TN_NAVY

    doc.add_paragraph("─" * 60)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(template_title)
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = TN_GOLD

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if is_english:
        meta_items = [
            f"Ref No: {ref_number}",
            f"Date: {date_str}",
            f"Officer ID: {officer_id}",
        ]
    else:
        meta_items = [
            f"குறிப்பு எண் / Ref No: {ref_number}",
            f"நாள் / Date: {date_str}",
            f"அலுவலர் / Officer: {officer_id}",
        ]
    for item in meta_items:
        run = meta_para.add_run(item + "\n")
        run.font.size = Pt(10)
        run.font.color.rgb = TN_NAVY

    doc.add_paragraph("─" * 60)


def _add_footer(doc: Document, is_english: bool = False):
    """Add government-styled footer / signature block."""
    doc.add_paragraph("─" * 60)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if is_english:
        lines = [
            "On behalf of the District Collector,",
            "Erode District.",
        ]
    else:
        lines = [
            "மாவட்ட ஆட்சியர் சார்பாக,",
            "ஈரோடு மாவட்டம்.",
        ]
    for line in lines:
        run = footer_para.add_run(line + "\n")
        run.font.size = Pt(10)
        run.font.color.rgb = TN_NAVY

    audit_para = doc.add_paragraph()
    audit_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    audit_run = audit_para.add_run(
        f"── AI-Assisted Official Draft | Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | "
        "Erode Collectorate AI System ──"
    )
    audit_run.font.size = Pt(7)
    audit_run.font.italic = True
    audit_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def _export_press_release_docx(doc: Document, content_data: Dict[str, Any], is_english: bool = False):
    """Format specifically according to DIPR Press Release layout (English / Tamil)."""
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    row = header_table.rows[0]

    # Left: Press Release Ref
    p_left = row.cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_label = f"PRESS RELEASE NO: {content_data['ref_number']}" if is_english else f"செ.வெ.எண் - {content_data['ref_number']}"
    r_ref = p_left.add_run(ref_label)
    r_ref.font.size = Pt(11)
    r_ref.font.bold = True
    r_ref.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # Right: Date
    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_label = f"DATE: {content_data['date_display']}" if is_english else f"நாள் - {content_data['date_display']}"
    r_date = p_right.add_run(date_label)
    r_date.font.size = Pt(11)
    r_date.font.bold = True
    r_date.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    title_p1 = doc.add_paragraph()
    title_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_english:
        title_text = "PRESS RELEASE ISSUED BY THE DISTRICT COLLECTOR & DISTRICT MAGISTRATE\nTHIRU S. KANDASAMY, I.A.S., ERODE DISTRICT"
    else:
        title_text = "ஈரோடு மாவட்ட ஆட்சித்தலைவர் திரு.ச.கந்தசாமி இ.ஆ.ப.,\nஅவர்களின் செய்திக்குறிப்பு-"
    r_title1 = title_p1.add_run(title_text)
    r_title1.font.size = Pt(12.5)
    r_title1.font.bold = True
    r_title1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = sep_p.add_run("----")
    r_sep.font.size = Pt(12)
    r_sep.font.bold = True
    r_sep.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    body_text = content_data.get("content_body", "")
    default_f = "Issued by: District Public Relations Officer, Erode District." if is_english else "வெளியீடு செய்தி மக்கள் தொடர்பு அலுவலர், ஈரோடு மாவட்டம்."
    clean_paras, final_footer = _extract_clean_body_paragraphs_and_footer(body_text, default_footer=default_f)

    for p_text in clean_paras:
        body_p = doc.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.line_spacing = 1.3
        body_p.paragraph_format.space_after = Pt(10)
        run = body_p.add_run(p_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph("─" * 60)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_footer = footer_p.add_run(final_footer)
    r_footer.font.size = Pt(10.5)
    r_footer.font.bold = True
    r_footer.font.color.rgb = RGBColor(0x00, 0x00, 0x00)


def _export_circular_docx(doc: Document, content_data: Dict[str, Any], is_english: bool = False):
    """Format specifically according to Erode District Collectorate Circular layout."""
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    row = header_table.rows[0]

    p_left = row.cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_label = f"CIRCULAR NO: {content_data['ref_number']}" if is_english else f"சுற்றறிக்கை எண் - {content_data['ref_number']}"
    r_ref = p_left.add_run(ref_label)
    r_ref.font.size = Pt(11)
    r_ref.font.bold = True
    r_ref.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_label = f"DATE: {content_data['date_display']}" if is_english else f"நாள் - {content_data['date_display']}"
    r_date = p_right.add_run(date_label)
    r_date.font.size = Pt(11)
    r_date.font.bold = True
    r_date.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    title_p1 = doc.add_paragraph()
    title_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_english:
        title_text = "OFFICE OF THE DISTRICT COLLECTOR, ERODE DISTRICT\nOFFICIAL CIRCULAR"
    else:
        title_text = "ஈரோடு மாவட்ட ஆட்சித்தலைவர் திரு.ச.கந்தசாமி இ.ஆ.ப.,\nஅவர்களின் சுற்றறிக்கை-"
    r_title1 = title_p1.add_run(title_text)
    r_title1.font.size = Pt(13)
    r_title1.font.bold = True
    r_title1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = sep_p.add_run("----")
    r_sep.font.size = Pt(12)
    r_sep.font.bold = True
    r_sep.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    body_text = content_data.get("content_body", "")
    default_f = "By Order of the District Collector, Erode District." if is_english else "மாவட்ட ஆட்சித்தலைவர் அவர்களின் உத்தரவுப்படி, ஈரோடு மாவட்டம்."
    clean_paras, final_footer = _extract_clean_body_paragraphs_and_footer(body_text, default_footer=default_f)

    for p_text in clean_paras:
        body_p = doc.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.line_spacing = 1.3
        body_p.paragraph_format.space_after = Pt(10)
        run = body_p.add_run(p_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph("─" * 60)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_footer = footer_p.add_run(final_footer)
    r_footer.font.size = Pt(10.5)
    r_footer.font.bold = True


def _export_memo_docx(doc: Document, content_data: Dict[str, Any], is_english: bool = False):
    """Format specifically according to Erode District Collectorate Office Memorandum layout."""
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    row = header_table.rows[0]

    p_left = row.cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_label = f"MEMORANDUM NO: {content_data['ref_number']}" if is_english else f"குறிப்பாணை எண் - {content_data['ref_number']}"
    r_ref = p_left.add_run(ref_label)
    r_ref.font.size = Pt(11)
    r_ref.font.bold = True
    r_ref.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_label = f"DATE: {content_data['date_display']}" if is_english else f"நாள் - {content_data['date_display']}"
    r_date = p_right.add_run(date_label)
    r_date.font.size = Pt(11)
    r_date.font.bold = True
    r_date.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph()

    title_p1 = doc.add_paragraph()
    title_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_english:
        title_text = "OFFICE OF THE DISTRICT COLLECTOR, ERODE DISTRICT\nOFFICE MEMORANDUM"
    else:
        title_text = "ஈரோடு மாவட்ட ஆட்சித்தலைவர் திரு.ச.கந்தசாமி இ.ஆ.ப.,\nஅவர்களின் அலுவலகக் குறிப்பாணை-"
    r_title1 = title_p1.add_run(title_text)
    r_title1.font.size = Pt(13)
    r_title1.font.bold = True
    r_title1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = sep_p.add_run("----")
    r_sep.font.size = Pt(12)
    r_sep.font.bold = True
    r_sep.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    body_text = content_data.get("content_body", "")
    default_f = "Personal Assistant to the District Collector, Erode District." if is_english else "மாவட்ட ஆட்சித்தலைவர் அவர்களின் உத்தரவுப்படி, ஈரோடு மாவட்டம்."
    clean_paras, final_footer = _extract_clean_body_paragraphs_and_footer(body_text, default_footer=default_f)

    for p_text in clean_paras:
        body_p = doc.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.line_spacing = 1.3
        body_p.paragraph_format.space_after = Pt(10)
        run = body_p.add_run(p_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph("─" * 60)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_footer = footer_p.add_run(final_footer)
    r_footer.font.size = Pt(10.5)
    r_footer.font.bold = True


def _export_meeting_minutes_docx(doc: Document, content_data: Dict[str, Any], is_english: bool = False):
    """Format specifically according to Erode District Meeting Minutes layout."""
    title_p1 = doc.add_paragraph()
    title_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_english:
        title_text = (
            f"PROCEEDINGS OF THE DISTRICT COLLECTOR & DISTRICT MAGISTRATE, ERODE\n"
            f"PRESENT: THIRU S. KANDASAMY, I.A.S., DISTRICT COLLECTOR\n"
            f"MINUTES OF REVIEW MEETING: {content_data['subject'].upper()}"
        )
    else:
        title_text = (
            f"ஈரோடு மாவட்ட ஆட்சித்தலைவர் அவர்கள் தலைமையில் {content_data['date_display']} அன்று\n"
            f"நடைபெற்ற {content_data['subject']} கூட்ட நடவடிக்கைகள்\n"
            f"முன்னிலை: திரு.ச.கந்தசாமி, இ.ஆ.ப.,"
        )
    r_title1 = title_p1.add_run(title_text)
    r_title1.font.size = Pt(12)
    r_title1.font.bold = True
    r_title1.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = True
    row = header_table.rows[0]

    p_left = row.cells[0].paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ref_label = f"Roc. No: {content_data['ref_number']}/2026" if is_english else f"எண்: வே/{content_data['ref_number']}/2026"
    r_ref = p_left.add_run(ref_label)
    r_ref.font.size = Pt(11)
    r_ref.font.bold = True
    r_ref.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    p_right = row.cells[1].paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_label = f"Dated: {content_data['date_display']}" if is_english else f"நாள்: {content_data['date_display']}"
    r_date = p_right.add_run(date_label)
    r_date.font.size = Pt(11)
    r_date.font.bold = True
    r_date.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph("─" * 60)

    subj_p = doc.add_paragraph()
    sub_label = f"Sub: {content_data['subject']} – Minutes of Review Meeting – Approval and Orders – Issued." if is_english else f"பொருள்: {content_data['subject']} – கூட்ட நடவடிக்கைகள் – ஒப்புதல் அளித்தல் – தொடர்பாக."
    r_subj = subj_p.add_run(sub_label)
    r_subj.font.size = Pt(10.5)
    r_subj.font.bold = True
    r_subj.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    ref_p = doc.add_paragraph()
    ref_label = "Ref: G.O. Ms. No. 78, Agriculture & Farmers Welfare Department, dated 17.02.2016." if is_english else "பார்வை: அரசாணை எண்: 78 வேளாண்மை (வே.உ.6) துறை, நாள்: 17.02.2016."
    r_ref_line = ref_p.add_run(ref_label)
    r_ref_line.font.size = Pt(10)
    r_ref_line.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    sep_p = doc.add_paragraph()
    sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sep = sep_p.add_run("<><><>")
    r_sep.font.size = Pt(11)
    r_sep.font.bold = True

    body_text = content_data.get("content_body", "")
    clean_paras = _extract_clean_body_paragraphs(body_text)

    for p_text in clean_paras:
        body_p = doc.add_paragraph()
        body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_p.paragraph_format.line_spacing = 1.3
        body_p.paragraph_format.space_after = Pt(10)
        run = body_p.add_run(p_text)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    doc.add_paragraph("─" * 60)

    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if is_english:
        sign_text = "Sd/- S. Kandasamy\nDistrict Collector,\nErode District.\n"
    else:
        sign_text = "ஓம்/-ச.கந்தசாமி\nமாவட்ட ஆட்சித்தலைவர்,\nஈரோடு.\n"
    r_sign = sign_p.add_run(sign_text)
    r_sign.font.size = Pt(11)
    r_sign.font.bold = True
    r_sign.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    pa_p = doc.add_paragraph()
    pa_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_english:
        pa_text = "/ By Order /\n\nPersonal Assistant to District Collector,\nDistrict Collectorate,\nErode."
    else:
        pa_text = "/உத்தரவுப்படி/\n\nநேர்முக உதவியாளர்,\nமாவட்ட ஆட்சியர் அலுவலகம்,\nஈரோடு."
    r_pa = pa_p.add_run(pa_text)
    r_pa.font.size = Pt(10)
    r_pa.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def export_to_docx(content_data: Dict[str, Any]) -> Path:
    """Export generated content to a formatted .docx file with language awareness."""
    doc = Document()

    # Set margins (0.8 inch)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    template_type = content_data.get("template_type", "press_release")
    lang = content_data.get("language") or detect_language(content_data.get("subject", "") + " " + content_data.get("content_body", ""))
    is_english = (lang == "en")

    # Route to specialized authentic layout
    if template_type == "press_release":
        _export_press_release_docx(doc, content_data, is_english=is_english)
    elif template_type == "circular":
        _export_circular_docx(doc, content_data, is_english=is_english)
    elif template_type == "memo":
        _export_memo_docx(doc, content_data, is_english=is_english)
    elif template_type == "meeting_minutes":
        _export_meeting_minutes_docx(doc, content_data, is_english=is_english)
    else:
        # Generic official document
        title = content_data.get("template_title_en" if is_english else "template_title_ta", "Official Document")
        _add_header(
            doc,
            template_title=title,
            ref_number=content_data.get("ref_number", "ERD/GEN/2026/001"),
            date_str=content_data.get("date_display", datetime.now().strftime("%d.%m.%Y")),
            officer_id=content_data.get("officer_id", "OFC001"),
            is_english=is_english,
        )
        body_text = content_data.get("content_body", "")
        clean_paras = _extract_clean_body_paragraphs(body_text)
        for p_text in clean_paras:
            body_p = doc.add_paragraph()
            body_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body_p.paragraph_format.line_spacing = 1.3
            body_p.paragraph_format.space_after = Pt(10)
            run = body_p.add_run(p_text)
            run.font.size = Pt(11)
        _add_footer(doc, is_english=is_english)

    output_dir = config.OUTPUTS_CONTENT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{content_data['content_id']}_{content_data['template_type']}.docx"
    output_path = output_dir / filename

    doc.save(str(output_path))
    logger.info(f"DOCX exported: {output_path}")
    return output_path


def export_to_pdf(content_data: Dict[str, Any]) -> Path:
    """Export generated content directly to an official formatted .pdf file."""
    import pymupdf

    font_path = "C:/Windows/Fonts/Nirmala.ttc"
    if not Path(font_path).exists():
        font_path = None

    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)  # A4

    c_navy = (0.10, 0.23, 0.36)
    c_dark = (0.13, 0.13, 0.13)
    c_gray = (0.50, 0.50, 0.50)

    template_type = content_data.get("template_type", "press_release")
    lang = content_data.get("language") or detect_language(content_data.get("subject", "") + " " + content_data.get("content_body", ""))
    is_english = (lang == "en")

    if template_type == "press_release":
        ref_text = f"PRESS RELEASE NO: {content_data['ref_number']}" if is_english else f"செ.வெ.எண் - {content_data['ref_number']}"
        date_text = f"DATE: {content_data['date_display']}" if is_english else f"நாள் - {content_data['date_display']}"
        header_text = "GOVERNMENT OF TAMIL NADU — ERODE DISTRICT" if is_english else "ஈரோடு மாவட்டம்"
        footer_text = "Issued by: District Public Relations Officer, Erode District." if is_english else "வெளியீடு - செய்தி மக்கள் தொடர்பு அலுவலர், ஈரோடு மாவட்டம்."

        page.insert_text((50, 55), ref_text, fontfile=font_path, fontsize=10.5, color=c_navy)
        page.insert_text((410, 55), date_text, fontfile=font_path, fontsize=10.5, color=c_navy)
        page.insert_text((180 if is_english else 235, 95), header_text, fontfile=font_path, fontsize=13, color=c_navy)

        subj_rect = pymupdf.Rect(50, 110, 545, 160)
        page.insert_textbox(subj_rect, content_data['subject'], fontfile=font_path, fontsize=11.5, color=c_navy, align=1)
        page.draw_line(pymupdf.Point(50, 165), pymupdf.Point(545, 165), color=c_navy, width=0.8)

        body_text = content_data.get("content_body", "")
        cleaned_paras = _extract_clean_body_paragraphs(body_text)
        full_body = "\n\n".join(cleaned_paras)

        body_rect = pymupdf.Rect(50, 180, 545, 750)
        page.insert_textbox(body_rect, full_body, fontfile=font_path, fontsize=10.5, color=c_dark, lineheight=1.5)

        page.draw_line(pymupdf.Point(50, 770), pymupdf.Point(545, 770), color=c_gray, width=0.5)
        page.insert_text((50, 790), footer_text, fontfile=font_path, fontsize=9.5, color=c_navy)

    else:
        top_hdr = "GOVERNMENT OF TAMIL NADU — DISTRICT COLLECTORATE" if is_english else "தமிழ்நாடு அரசு — மாவட்ட ஆட்சியர் அலுவலகம்"
        dist_hdr = "ERODE DISTRICT" if is_english else "ஈரோடு மாவட்டம்"
        title_hdr = content_data.get("template_title_en", "Official Document") if is_english else f"{content_data.get('template_title_ta', '')} ({content_data.get('template_title_en', '')})"
        ref_text = f"Ref No: {content_data['ref_number']}" if is_english else f"குறிப்பு எண்: {content_data['ref_number']}"
        date_text = f"Date: {content_data['date_display']}" if is_english else f"நாள்: {content_data['date_display']}"

        page.insert_text((150 if is_english else 180, 55), top_hdr, fontfile=font_path, fontsize=12, color=c_navy)
        page.insert_text((235 if is_english else 245, 75), dist_hdr, fontfile=font_path, fontsize=11, color=c_navy)
        page.insert_text((200, 105), title_hdr, fontfile=font_path, fontsize=12, color=c_navy)

        page.insert_text((50, 130), ref_text, fontfile=font_path, fontsize=10, color=c_navy)
        page.insert_text((430, 130), date_text, fontfile=font_path, fontsize=10, color=c_navy)

        page.draw_line(pymupdf.Point(50, 145), pymupdf.Point(545, 145), color=c_navy, width=0.8)

        body_text = content_data.get("content_body", "")
        cleaned_paras = _extract_clean_body_paragraphs(body_text)
        full_body = "\n\n".join(cleaned_paras)

        body_rect = pymupdf.Rect(50, 160, 545, 740)
        page.insert_textbox(body_rect, full_body, fontfile=font_path, fontsize=10.5, color=c_dark, lineheight=1.5)

        page.draw_line(pymupdf.Point(50, 755), pymupdf.Point(545, 755), color=c_gray, width=0.5)
        sign_l1 = "On behalf of the District Collector," if is_english else "மாவட்ட ஆட்சியர் சார்பாக,"
        sign_l2 = "Erode District." if is_english else "ஈரோடு மாவட்டம்."
        page.insert_text((360 if is_english else 380, 775), sign_l1, fontfile=font_path, fontsize=10, color=c_navy)
        page.insert_text((400 if is_english else 410, 790), sign_l2, fontfile=font_path, fontsize=10, color=c_navy)

    output_dir = config.OUTPUTS_CONTENT_DIR
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{content_data['content_id']}_{content_data['template_type']}.pdf"
    output_path = output_dir / filename
    doc.save(str(output_path))
    logger.info(f"PDF exported: {output_path}")
    return output_path
