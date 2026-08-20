"""Anti-Hallucination Grounded Draft Generation Layer & DOCX Export."""

import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import jinja2

import config
from pipeline.database import log_audit, save_draft, update_source_status

logger = logging.getLogger("TamilDraftGenerator")


def export_draft_to_docx(draft_text: str, source_id: str, output_path: Path) -> Path:
    """Generate official Tamil Government acknowledgment letter as a styled .docx document."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = docx.Document()

    # Configure 1-inch standard government margins
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Style: Normal font configuration
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Nirmala UI"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    # Header Title Banner
    h_p = doc.add_paragraph()
    h_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h_run1 = h_p.add_run("தமிழ்நாடு அரசு\n")
    h_run1.bold = True
    h_run1.font.size = Pt(14)
    h_run1.font.color.rgb = RGBColor(0x1E, 0x3C, 0x72)

    h_run2 = h_p.add_run("மாவட்ட ஆட்சியர் அலுவலகம், ஈரோடு மாவட்டம்\n")
    h_run2.bold = True
    h_run2.font.size = Pt(12)

    h_run3 = h_p.add_run("மக்கள் குறைதீர்க்கும் மனு ஒப்புகைச் சீட்டு (Acknowledgment Slip)\n")
    h_run3.italic = True
    h_run3.font.size = Pt(10)
    h_run3.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph("―" * 45).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Content body
    lines = draft_text.splitlines()
    for line in lines:
        line_str = line.strip()
        if not line_str or line_str.startswith("===") or line_str.startswith("---") or "தமிழ்நாடு அரசு" in line_str or "மாவட்ட ஆட்சியர் அலுவலகம்" in line_str:
            continue

        p = doc.add_paragraph()
        if ":" in line_str and not line_str.startswith("விபரம்") and not line_str.startswith("பொருள்"):
            parts = line_str.split(":", 1)
            lbl_run = p.add_run(parts[0] + ": ")
            lbl_run.bold = True
            
            val_str = parts[1].strip()
            val_run = p.add_run(val_str)
            if "தகவல் இல்லை" in val_str:
                val_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                val_run.italic = True
        else:
            p.add_run(line_str)

    # Footer note
    doc.add_paragraph("\n")
    f_p = doc.add_paragraph()
    f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_run = f_p.add_run("மாவட்ட ஆட்சியரின் நேர்முக உதவியாளர்,\nஈரோடு மாவட்டம்.")
    f_run.bold = True

    # Provenance meta footer
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_run = meta_p.add_run(f"\n[System Document ID: {source_id}]")
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


class TamilDraftGenerator:
    """Renders deterministic Tamil acknowledgment drafts and enforces strict anti-hallucination contracts."""

    def __init__(self, templates_dir: Optional[Path] = None):
        self.templates_dir = templates_dir or config.TEMPLATES_DIR
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(self.templates_dir)),
            autoescape=False,
            undefined=jinja2.DebugUndefined,
        )

    def select_template(self, department: str) -> Tuple[str, str]:
        """Select appropriate Jinja2 template file for department."""
        if department in ("வருவாய்", "பதிவுத்துறை"):
            return "ack_revenue_tamil.txt", "Revenue/Registration Acknowledgment"
        elif department == "சமூக_நலன்":
            return "ack_social_tamil.txt", "Social Welfare Acknowledgment"
        else:
            return "ack_general_tamil.txt", "General Grievance Acknowledgment"

    def render_draft(
        self,
        source_id: str,
        department: str,
        extracted_entities: Dict[str, Any],
    ) -> Dict[str, Any]:
        """Render draft with deterministic substitution, missing slots detection, and grounding map persistence."""
        template_filename, template_desc = self.select_template(department)

        expected_slots = [
            "file_number",
            "date",
            "applicant_name",
            "mobile_number",
            "aadhaar_number",
            "taluk",
            "village",
        ]
        if department == "வருவாய்":
            expected_slots.append("survey_number")

        sanitized_context = {}
        missing_fields: List[str] = []
        grounding_map: Dict[str, Any] = extracted_entities.get("_grounding_map", {})

        for slot in expected_slots:
            val = extracted_entities.get(slot)
            if not val or val == config.MISSING_DATA_PLACEHOLDER or str(val).strip() == "":
                sanitized_context[slot] = f"{config.MISSING_DATA_PLACEHOLDER} — கைமுறையாக நிரப்பவும்"
                missing_fields.append(slot)
                if slot not in grounding_map:
                    grounding_map[slot] = {
                        "value": None,
                        "source": None,
                        "confidence": 0.0,
                        "validation_status": "missing",
                        "source_chunk": "",
                    }
            else:
                sanitized_context[slot] = str(val).strip()

        sanitized_context["department"] = department
        sanitized_context["grievance_summary"] = extracted_entities.get(
            "grievance_summary", "மனு பரிசீலனை கோருதல்"
        )
        if "survey_number" not in sanitized_context:
            sanitized_context["survey_number"] = extracted_entities.get(
                "survey_number", f"{config.MISSING_DATA_PLACEHOLDER} — கைமுறையாக நிரப்பவும்"
            )

        # Compute hallucination / incomplete metric
        hallucination_score = round(len(missing_fields) / max(1, len(expected_slots)), 3)

        # Render template
        template = self.jinja_env.get_template(template_filename)
        draft_text = template.render(**sanitized_context)

        # Persist draft with grounding map and missing fields list
        save_draft(
            source_id=source_id,
            draft_text=draft_text,
            template_used=template_filename,
            hallucination_score=hallucination_score,
            grounding_map=grounding_map,
            missing_fields=missing_fields,
        )
        update_source_status(source_id=source_id, status="draft_ready")
        log_audit(
            source_id=source_id,
            action="DRAFT_GENERATED",
            officer_id="SYSTEM_DRAFTER",
            details=f"Rendered {template_filename} with {len(missing_fields)} missing fields (hallucination_score: {hallucination_score})",
        )

        return {
            "source_id": source_id,
            "draft_text": draft_text,
            "template_used": template_filename,
            "hallucination_score": hallucination_score,
            "missing_fields": missing_fields,
            "grounding_map": grounding_map,
        }
